from flask import Flask, render_template, request, send_file
from docx import Document
import requests
import os

app = Flask(__name__)

GROQ_API_KEY = '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'
SERPER_API_KEY = 'f99c6211fd876adefce7761dda3b36eea33980f7'
GROQ_API_URL = 'https://api.groq.com/openai/v1/chat/completions'
SERPER_API_URL = 'https://google.serper.dev/search'

def query_groq_api(prompt):
    headers = {
        'Authorization': f'Bearer {GROQ_API_KEY}',
        'Content-Type': 'application/json',
    }
    data = {
        'model': 'llama3-8b-8192',
        'messages': [{'role': 'user', 'content': str(prompt)}],
    }
    response = requests.post(GROQ_API_URL, headers=headers, json=data)
    response.raise_for_status()
    return response.json()

def extract_text_from_groq_response(response):
    if 'choices' in response and len(response['choices']) > 0:
        return response['choices'][0]['message']['content']
    return "Error: Unexpected response format from Groq API."

def create_legal_contract_from_text(text):
    document = Document()
    document.add_heading('Lease Agreement', level=1)
    document.add_paragraph(text)
    document_path = 'contract.docx'
    document.save(document_path)
    return document_path

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/generate', methods=['POST'])
def generate():
    user_input = request.form['user_input']
    response = query_groq_api(user_input)
    lease_agreement_text = extract_text_from_groq_response(response)
    document_path = create_legal_contract_from_text(lease_agreement_text)
    return send_file(document_path, as_attachment=True)

if __name__ == '__main__':
    app.run(debug=True)
